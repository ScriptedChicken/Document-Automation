import pandas as pd
from docx import Document
from datetime import datetime as dt
import os


def populate_table(df, table_i):
    df_filtered = df.loc[:, df.columns != 'Invoice Number']
    table = document.tables[table_i]
    for i in range(df_filtered.shape[0]):
        for j in range(df_filtered.shape[-1]):
            value = df_filtered.values[i, j]
            if type(value) == float:
                value = round(value, 2)
            table.cell(i+1, j).text = str(value)


def populate_transposed_table(df, table_i, padding=1):
    table = document.tables[table_i]
    for i in range(df.shape[-1]):
        for j in range(df.shape[0]):
            table.cell(j, i+padding).text = str(df.values[j, i])


def return_summary_df(item_data_df):
    df_filtered = item_data_df[item_data_df["Invoice Number"] == invoice_number]
    summary_df = pd.DataFrame(columns=['Subtotal', 'Sales Tax', 'Grand Total'])
    summary_df.loc[len(summary_df)] = [None, None, None]
    summary_df['Subtotal'] = df_filtered['Line Total'].sum()
    summary_df['Sales Tax'] = summary_df['Subtotal'] * 0.33
    summary_df['Grand Total'] = summary_df['Subtotal'] + summary_df['Sales Tax']
    summary_df = round(summary_df, 2)
    return summary_df.transpose()


customer_details_path = r"Inputs\Customer Details Multiple.xlsx"
item_data_path = r"Inputs\Item Data.xlsx"
word_template_path = r"Inputs\Example_Invoice.docx"
output_folder_path = r".\Outputs"

customer_details_df = pd.read_excel(customer_details_path)
item_data_df = pd.read_excel(item_data_path)

# add customer details
for cd_row_id, cd_row in customer_details_df.iterrows():
    document = Document(word_template_path)
    invoice_number = cd_row['Invoice Number']

    customer_details_df_filtered = customer_details_df[customer_details_df["Invoice Number"] == invoice_number]
    item_data_df_filtered = item_data_df[item_data_df["Invoice Number"] == invoice_number]

    # rename title to invoice number
    for paragraph in document.paragraphs:
        for i, run in enumerate(paragraph.runs):
            if "INVOICE" in run.text:
                paragraph.runs[i].text = f"INVOICE - {invoice_number}"

    # add recipient information
    recipient_information = customer_details_df_filtered[['Owner Name', 'Business Name', 'Business Address', 'Phone Number']]
    populate_transposed_table(recipient_information.transpose(), 0, padding=0)

    # add customer details
    customer_details_df_filtered = customer_details_df_filtered[['Salesperson', 'Job Title', 'Payment Terms', 'Due Date', 'Invoice Number']]
    customer_details_df_filtered['Due Date'] = customer_details_df_filtered['Due Date'].dt.strftime('%d/%m/%Y')
    populate_table(customer_details_df_filtered, 2)

    # add item details
    populate_table(item_data_df_filtered, 3)

    # add subtotal, tax, and grand total
    summary_df = return_summary_df(item_data_df_filtered)
    populate_transposed_table(summary_df, 4)

    # save
    new_docx_path = os.path.join('Outputs', f"INVOICE_{invoice_number}.docx")
    document.save(new_docx_path)

print("Complete")
